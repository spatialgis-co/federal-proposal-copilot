#!/usr/bin/env python3
"""
tools/brand_polish.py — Apply SpatialGIS brand styling to exported .docx files.

Reads my-company/brand-palette.md (Spatial Blue #4E74B9, body text #2D2D2D, etc.)
and applies them to heading styles, table headers, and document footer.

This is a separate post-processing step (not part of md_to_docx.py) so that:
  - Non-SpatialGIS users of the framework get unbranded output by default
  - Brand changes are isolated and reversible
  - The same export can be brand-polished or not on demand

Usage:
    python tools/brand_polish.py --proposal va-gis-rfi
    python tools/brand_polish.py --files path/to/a.docx path/to/b.docx
    python tools/brand_polish.py --proposal va-gis-rfi --no-footer  # skip footer rewrite

Idempotent: re-running re-applies the same polish. If md_to_docx.py is re-run,
re-apply brand_polish.py afterwards (the export step does not preserve it).
"""
import argparse
import glob
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx not installed. Run: pip install python-docx\n")
    sys.exit(1)

# ── SpatialGIS brand palette (from my-company/brand-palette.md) ──────────────
SPATIAL_BLUE = RGBColor(0x4E, 0x74, 0xB9)
SPATIAL_GREEN = RGBColor(0xA8, 0xD5, 0x9C)
TEXT = RGBColor(0x2D, 0x2D, 0x2D)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GRAY = "F5F5F5"  # hex for table shading

COMPANY_FOOTER = "SpatialGIS, LLC  |  CAGE: 7RFJ7  |  UEI: XE8LEMK77DC9"


# ── Style helpers ────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    """Apply a background shading color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def style_headings(doc):
    """Set heading colors to Spatial Blue."""
    heading_specs = [
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11, True),
        ("Heading 4", 11, True),
    ]
    for name, size, bold in heading_specs:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        font = style.font
        font.color.rgb = SPATIAL_BLUE
        font.bold = bold
        font.size = Pt(size)


def style_body(doc):
    """Set body text to dark gray (#2D2D2D)."""
    try:
        normal = doc.styles["Normal"]
        normal.font.color.rgb = TEXT
    except KeyError:
        pass


def style_tables(doc):
    """Spatial Blue header row with white text; alternating row shading on body."""
    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        # Header row: Spatial Blue fill + white text + bold
        header = table.rows[0]
        for cell in header.cells:
            shade_cell(cell, "4E74B9")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
        # Alternating row shading on body rows (every other row)
        for i, row in enumerate(table.rows[1:], start=1):
            if i % 2 == 0:
                for cell in row.cells:
                    shade_cell(cell, LIGHT_GRAY)


def apply_footer(doc, footer_text=COMPANY_FOOTER):
    """Put the company footer (with CAGE/UEI) on every page."""
    for section in doc.sections:
        footer = section.footer
        # Clear any existing footer paragraphs (idempotent)
        for para in list(footer.paragraphs):
            para.clear()
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


# ── Main ─────────────────────────────────────────────────────────────────────

def polish(path: Path, skip_footer: bool = False) -> None:
    doc = Document(str(path))
    style_headings(doc)
    style_body(doc)
    style_tables(doc)
    if not skip_footer:
        apply_footer(doc)
    doc.save(str(path))


def main():
    ap = argparse.ArgumentParser(description="SpatialGIS brand polish for exported .docx")
    ap.add_argument("--proposal", help="Proposal slug under proposals/")
    ap.add_argument("--files", nargs="+", help="Explicit list of .docx files")
    ap.add_argument("--workspace", default=None, help="Workspace root (default: cwd or script-parent)")
    ap.add_argument("--no-footer", action="store_true", help="Skip footer rewrite")
    args = ap.parse_args()

    if not args.proposal and not args.files:
        ap.error("Must specify --proposal or --files")

    targets = []
    if args.proposal:
        workspace = Path(args.workspace) if args.workspace else Path(__file__).resolve().parent.parent
        docx_dir = workspace / "proposals" / args.proposal / "final" / "docx"
        if not docx_dir.exists():
            sys.exit(f"No final/docx dir: {docx_dir}")
        targets.extend(sorted(docx_dir.glob("*.docx")))
    if args.files:
        targets.extend(Path(p) for p in args.files)

    if not targets:
        sys.exit("No .docx files to polish")

    for t in targets:
        polish(t, skip_footer=args.no_footer)
        print(f"  [OK] {t.name}")
    print(f"Done — polished {len(targets)} file(s)")


if __name__ == "__main__":
    main()
